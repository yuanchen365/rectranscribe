# modules/doc_generator.py
"""
會議摘要報告 Word 文件產出模組
- 提供 generate_docx() 函式，供 batch_job.py 呼叫
- 可從 CLI 輸入 JSON 測試產出 Word 報告
"""

import os
import json
from docx import Document
from typing import List

def generate_docx(file_path: str, summary: str, outline: List[str], todos: List[str]) -> None:
    """
    將分析結果匯出為 Word 檔（.docx）

    Args:
        file_path (str): 輸出路徑
        summary (str): 總體摘要
        outline (List[str]): 議題大綱
        todos (List[str]): 待辦清單
    """
    doc = Document()
    doc.add_heading("會議摘要報告", level=0)

    doc.add_heading("摘要", level=1)
    doc.add_paragraph(summary.strip() if summary else "（無內容）")

    doc.add_heading("大綱", level=1)
    if outline:
        for item in outline:
            doc.add_paragraph(item.strip(), style="List Bullet")
    else:
        doc.add_paragraph("（無內容）")

    doc.add_heading("待辦事項", level=1)
    if todos:
        for item in todos:
            doc.add_paragraph(item.strip(), style="List Number")
    else:
        doc.add_paragraph("（無內容）")

    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)
    print(f"✅ 已輸出 Word 報告至：{file_path}")

# ✅ CLI 測試模式
if __name__ == "__main__":
    json_path = "output/final/meeting_summary.json"
    output_docx_path = "output/final/meeting_summary.docx"

    if not os.path.exists(json_path):
        print(f"❌ 找不到分析結果 JSON：{json_path}")
        exit(1)

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    summary = data.get("summary", "")
    outline = data.get("outline", [])
    todos = data.get("todos", [])

    generate_docx(output_docx_path, summary, outline, todos)
