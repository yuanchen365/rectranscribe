# -*- coding: utf-8 -*-
"""
批次處理（逐段驗證 → 合併 → 產出最終檔）
修正版：鎖定「本次 job 的分段資料夾」，避免誤抓根目錄舊檔。
"""
import os
import re
import json
import time
import hashlib
from pathlib import Path
from typing import List, Tuple, Dict, Any, Optional, Callable

# ✅ 修正 import 為 modules.xx
from modules.transcribe import transcribe
from modules.analyze import run_analysis
from modules.ai_precheck import ai_precheck  # ✅ 確保這個模組放在 modules/
from docx import Document
from modules.doc_generator import generate_docx  # 若你有另一版客製 DOCX，可選用

# 預設資料夾（根）
SEG_AUDIO_ROOT = Path("output") / "segments"
SEG_TEXT_ROOT  = Path("output") / "segments_text"
FINAL_ROOT     = Path("output") / "final"
LOGS_ROOT      = Path("output") / "logs"

for p in [SEG_AUDIO_ROOT, SEG_TEXT_ROOT, FINAL_ROOT, LOGS_ROOT]:
    p.mkdir(parents=True, exist_ok=True)


def _nkey(s: str):
    return [int(t) if t.isdigit() else t for t in re.split(r"(\d+)", s)]


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _save_text(path: Path, text: str):
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(text or "")


def _save_json(path: Path, data: Dict[str, Any]):
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _save_docx(path: Path, summary: str, outline: List[str], todos: List[str]):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("會議摘要報告", level=0)

    doc.add_heading("摘要", level=1)
    doc.add_paragraph(summary if summary else "（無內容）")

    doc.add_heading("大綱", level=1)
    if outline:
        for item in outline:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("（無內容）")

    doc.add_heading("待辦事項", level=1)
    if todos:
        for t in todos:
            doc.add_paragraph(t, style="List Number")
    else:
        doc.add_paragraph("（無內容）")

    doc.save(str(path))


def _list_segments(seg_dir: Path) -> List[Path]:
    files = [seg_dir / f for f in os.listdir(seg_dir)
             if (seg_dir / f).is_file() and (seg_dir / f).suffix.lower() in {".wav", ".mp3", ".m4a"}]
    files.sort(key=lambda p: _nkey(p.name))
    return files


_JOB_DIR_PAT = re.compile(r"^(job_|segments_)\d{8}_\d{6}$", re.IGNORECASE)


def _find_latest_job_dir(base: Path) -> Optional[Path]:
    """
    在 base 目錄下尋找符合 job_YYYYMMDD_HHMMSS 或 segments_YYYYMMDD_HHMMSS 的資料夾，
    以 mtime 最新者為目標。
    """
    candidates = []
    for child in base.iterdir():
        if child.is_dir() and _JOB_DIR_PAT.match(child.name):
            candidates.append(child)
    if not candidates:
        return None
    candidates.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return candidates[0]


def _job_id_from_dir(path: Path) -> str:
    """
    從資料夾名稱推導 job_id；若不符合慣例則用 timestamp。
    """
    name = path.name
    if _JOB_DIR_PAT.match(name):
        return name
    return "job_" + time.strftime("%Y%m%d_%H%M%S")


def run_batch_process(
    segments_dir: str = str(SEG_AUDIO_ROOT),
    preview: bool = False,
    force: bool = False,
    return_progress: bool = False,
    *,
    start_index: int = 1,
    max_segments: Optional[int] = None,
    on_progress: Optional[Callable[[int, int, str], None]] = None,
) -> str:
    """
    segments_dir:
        - 建議傳入「本次 job 的資料夾」(例如 output/segments/job_20251001_123009)
        - 若傳入的是根目錄 output/segments：
            * 若偵測到底下有 job_* 子資料夾，會自動改抓「最新的 job」(安全柵欄)
            * 若沒有 job 子資料夾，才會處理根目錄的檔案（不建議）
    """
    t0 = time.time()
    progress: List[str] = []

    def log(msg: str):
        s = time.strftime("%H:%M:%S") + " " + msg
        print(s)
        progress.append(s)

    seg_dir = Path(segments_dir).resolve()

    # ✅ 安全柵欄：如果傳的是根目錄，且底下存在 job_*，自動切換到最新 job
    if seg_dir.samefile(SEG_AUDIO_ROOT.resolve()):
        latest = _find_latest_job_dir(seg_dir)
        if latest:
            log(f"ℹ️ 你傳入的是根目錄：{seg_dir}")
            log(f"➡️  自動鎖定最新 job 目錄：{latest.name}")
            seg_dir = latest

    if not seg_dir.exists() or not seg_dir.is_dir():
        raise FileNotFoundError(f"找不到分段目錄：{seg_dir}")

    job_id = _job_id_from_dir(seg_dir)

    # 🔒 本次 job 的輸出路徑（避免互相覆蓋）
    seg_text_dir = (SEG_TEXT_ROOT / job_id)
    final_dir    = (FINAL_ROOT / job_id)
    logs_dir     = (LOGS_ROOT / job_id)

    seg_text_dir.mkdir(parents=True, exist_ok=True)
    final_dir.mkdir(parents=True, exist_ok=True)
    logs_dir.mkdir(parents=True, exist_ok=True)

    # 列出本次段檔
    all_seg_files = _list_segments(seg_dir)
    if not all_seg_files:
        # 再提供一次友善提示：是否把檔案放錯地方？
        root_examples = _list_segments(SEG_AUDIO_ROOT)
        hint = ""
        if root_examples:
            hint = f"；注意：在根目錄 {SEG_AUDIO_ROOT} 發現 {len(root_examples)} 個音檔，請確認是否放錯位置"
        raise FileNotFoundError(f"在 {seg_dir} 沒找到任何音檔{hint}。")

    start = max(0, start_index - 1)
    end = None if not max_segments or max_segments <= 0 else start + max_segments
    seg_files = all_seg_files[start:end]
    total = len(seg_files)

    log(f"🔧 Job: {job_id}")
    log(f"📂 Segments 來源：{seg_dir}")
    log(f"📦 文字輸出：{seg_text_dir}")
    log(f"📦 最終輸出：{final_dir}")
    log(f"🔧 將處理段落：從第 {start_index} 段開始；最多處理 {max_segments if max_segments else '全部'} 段。")
    log(f"📊 實際段數：{total}（全部共有 {len(all_seg_files)} 段）")

    manifest = {
        "job_id": job_id,
        "segments_dir": str(seg_dir),
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "preview": preview,
        "items": [],
        "start_index": start_index,
        "count": total
    }

    if on_progress:
        try:
            on_progress(0, total, "開始處理")
        except Exception:
            pass

    for idx, audio_path in enumerate(seg_files, 1):
        abs_index = start_index + idx - 1
        seg_id = f"{abs_index:02d}"
        base = f"seg_{seg_id}"
        paths = {
            "transcript": seg_text_dir / f"{base}_transcript.txt",
            "review":     seg_text_dir / f"{base}_review.txt",
            "revised":    seg_text_dir / f"{base}_revised.txt",
            "meta":       seg_text_dir / f"{base}_meta.json",
        }

        if (not force) and all(p.exists() for p in paths.values()):
            log(f"🟡 跳過 {base}（已有輸出）")
            try:
                with open(paths["meta"], "r", encoding="utf-8") as f:
                    meta = json.load(f)
            except Exception:
                meta = {
                    "seg_id": seg_id,
                    "audio": Path(audio_path).name,
                    "paths": {k: str(v) for k, v in paths.items()},
                    "ok": True,
                }
            manifest["items"].append(meta)
            if on_progress:
                try:
                    on_progress(idx, total, f"跳過第 {abs_index} 段")
                except Exception:
                    pass
            continue

        log(f"🎧 [{seg_id}] 處理：{Path(audio_path).name}")
        try:
            transcript = transcribe(str(audio_path))
            review_text, revised_text = ai_precheck(transcript, preview=preview)

            _save_text(paths["transcript"], transcript)
            _save_text(paths["review"], review_text)
            _save_text(paths["revised"], revised_text)

            meta = {
                "seg_id": seg_id,
                "audio": Path(audio_path).name,
                "audio_sha256": _sha256(Path(audio_path)),
                "paths": {k: str(v) for k, v in paths.items()},
                "lens": {
                    "transcript": len(transcript or ""),
                    "review": len(review_text or ""),
                    "revised": len(revised_text or "")
                },
                "ok": True,
            }
            _save_json(paths["meta"], meta)
            manifest["items"].append(meta)
            log(f"✅ [{seg_id}] 完成（T{meta['lens']['transcript']}/R{meta['lens']['revised']}）")
        except Exception as e:
            meta = {
                "seg_id": seg_id,
                "audio": Path(audio_path).name,
                "paths": {k: str(v) for k, v in paths.items()},
                "ok": False,
                "error": repr(e)
            }
            _save_json(paths["meta"], meta)
            manifest["items"].append(meta)
            log(f"❌ [{seg_id}] 失敗：{e!r}")

        if on_progress:
            try:
                on_progress(idx, total, f"完成第 {abs_index} 段")
            except Exception:
                pass

    manifest_path = logs_dir / f"batch_manifest_{job_id}.json"
    _save_json(manifest_path, manifest)

    ok_items = [it for it in manifest["items"] if it.get("ok")]
    tr_paths = [Path(it["paths"]["transcript"]) for it in ok_items if it.get("paths", {}).get("transcript")]
    rvw_paths = [Path(it["paths"]["review"])     for it in ok_items if it.get("paths", {}).get("review")]
    rvd_paths = [Path(it["paths"]["revised"])    for it in ok_items if it.get("paths", {}).get("revised")]

    def _stitch(paths: List[Path], header_prefix: str) -> str:
        chunks = []
        for i, p in enumerate(paths, 1):
            name = p.name
            with open(p, "r", encoding="utf-8") as f:
                txt = f.read()
            chunks.append(f"=== {header_prefix} {i:02d} | {name} ===\n{txt}\n")
        return "\n".join(chunks)

    merged_transcript = _stitch(tr_paths, "TRANSCRIPT SEG") if tr_paths else ""
    merged_review     = _stitch(rvw_paths, "REVIEW SEG")     if rvw_paths else ""
    merged_revised    = _stitch(rvd_paths, "REVISED SEG")    if rvd_paths else ""

    final_paths = {
        "transcript_txt":         final_dir / "transcript.txt",
        "transcript_review_txt":  final_dir / "transcript_review.txt",
        "transcript_revised_txt": final_dir / "transcript_revised.txt",
    }
    _save_text(final_paths["transcript_txt"], merged_transcript)
    _save_text(final_paths["transcript_review_txt"], merged_review)
    _save_text(final_paths["transcript_revised_txt"], merged_revised)

    # 分析與輸出報告
    summary, outline, todos = run_analysis(merged_revised, preview=False)

    meeting_paths = {
        "txt":  final_dir / "meeting_summary.txt",
        "docx": final_dir / "meeting_summary.docx",
        "json": final_dir / "meeting_summary.json",
    }
    _save_text(
        meeting_paths["txt"],
        "【摘要】\n" + (summary or "") +
        "\n\n【大綱】\n" + "\n".join(outline or []) +
        "\n\n【待辦事項】\n" + "\n".join(todos or [])
    )
    _save_docx(meeting_paths["docx"], summary, outline, todos)
    _save_json(meeting_paths["json"], {
        "summary": summary,
        "outline": outline,
        "todos": todos
    })

    log(f"📦 已產出報告：{meeting_paths['docx']}")
    log(f"🧾 Manifest：{manifest_path}")
    log(f"⏱️ 總耗時：{time.time() - t0:.1f} 秒")

    # ✅ 為 Flask app.py 回傳報告檔案路徑即可
    return str(meeting_paths["docx"])
