# modules/analyze.py
"""
功能升級（進階版摘要）：
1) 使用進階版 Prompt 產出「可決策摘要」：<背景/驅動>、<主要結論>、<建議行動>、<風險/前提>
2) 仍維持原介面：run_analysis(text, preview=False) -> (summary:str, outline:list[str], todos:list[str])
3) 模型以 JSON 回傳，解析更穩定；todos 支援物件/字串，最終轉為字串清單
4) 例外防護：模型失敗時回傳空值而不中斷流程
5) 可直接執行：python -m modules.analyze --preview
"""

from __future__ import annotations

import os
import sys
import json
from typing import Any, Dict, List, Tuple, cast

from dotenv import load_dotenv
from openai import OpenAI
from openai.types.chat import ChatCompletionMessageParam
from docx import Document

# === 初始化環境與 OpenAI 客戶端 ===
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
client = OpenAI(api_key=OPENAI_API_KEY)

# === 進階版 Prompt（可決策摘要；含評分規則與禁用） ===
ADVANCED_SUMMARY_PROMPT = """
角色：資深策略顧問（製造與供應鏈）
任務：把會議逐字稿濃縮為「可決策摘要」，供經營團隊5分鐘內讀懂並可排程執行。

【輸入資訊】
- 會議主題：{topic}
- 關鍵地理/政策：{regions_policies}
- 讀者：{audience}
- 風格：務實、可落地；避免行話。
- 最長字數（摘要）：{max_chars}

【輸出格式（請以 JSON 物件回覆；務必遵守鍵名）】
{{
  "summary": "string — 決策導向的摘要；請在 {max_chars} 字以內",
  "outline": ["string", "string", "string"],  // 主要議題清單（至少3點）
  "todos": [
    // 可以是字串，或下列物件；最終將被轉為字串顯示
    // {{ "action": "動作", "owner": "負責角色", "due": "時程或里程碑", "kpi": "衡量指標" }}
  ],
  "risks": ["string", "string"],               // 風險/前提（1–2條；可選）
  "self_evaluation": {{
    "score": 0-10,                             // 模型自評分（可選）
    "issues": ["缺少量化", "缺少風險", "超出字數"]  // 缺失項（可選）
  }}
}}

【內容規則】
- <背景/驅動>（1–2句）寫進 "summary" 的開頭；說明為何現在需要處理（政策/成本/風險/時程）
- <主要結論>（3點）放入 "outline"；每點需含「主詞 + 動作/影響 + 可驗證指標或等級（高/中/低/區間%）」
- <建議行動>（3條）放入 "todos"；若輸出為物件，請包含 action/owner/due/kpi
- <風險/前提>（1–2條）放入 "risks"；包含「監測指標 + 緩解手段」
- 禁用：空泛形容詞（大力、持續、積極）、AI 自述、逐字稿原句的直接貼上或冗長重複
- 若逐字稿有數據→直接用；沒有→以範圍/等級/里程碑表示（例如「良率T+6月≥95%」）

【逐字稿】
{transcript}
""".strip()


def _build_advanced_messages(
    transcript: str,
    *,
    topic: str = "供應鏈重構與關稅因應",
    regions_policies: str = "美國關稅、CBAM、印度/墨西哥友岸化",
    audience: str = "經營團隊",
    max_chars: int = 220,
) -> List[ChatCompletionMessageParam]:
    prompt = ADVANCED_SUMMARY_PROMPT.format(
        topic=topic,
        regions_policies=regions_policies,
        audience=audience,
        max_chars=max_chars,
        transcript=transcript.strip(),
    )
    messages: List[ChatCompletionMessageParam] = [
        {
            "role": "system",
            "content": "你是嚴謹的策略顧問，擅長把冗長會議轉為可執行摘要與行動清單。"
        },
        {"role": "user", "content": prompt},
    ]
    return messages


def ai_analyze(text: str) -> str:
    """
    呼叫 GPT（進階版）進行分析，回傳 JSON 字串。
    與舊版同名以保持相容；內部已升級為進階版 prompt。
    """
    try:
        messages = _build_advanced_messages(
            text,
            topic="供應鏈重構與關稅因應",
            regions_policies="美國關稅、CBAM、印度/墨西哥友岸化",
            audience="經營團隊",
            max_chars=220,
        )
        typed_messages = cast(List[ChatCompletionMessageParam], messages)
        resp = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=typed_messages,
            temperature=0.3,
            presence_penalty=0.0,
            frequency_penalty=0.2,
            response_format={"type": "json_object"},
            max_tokens=900,
        )
        return resp.choices[0].message.content or ""
    except Exception as e:
        print(f"❌ OpenAI 分析錯誤：{e}")
        return "{}"  # 不中斷流程


def _stringify_todo(item: Any) -> str:
    """
    將 todos 的元素（可能是字串或物件）轉為乾淨的字串。
    物件鍵支援：action/owner/due/kpi；其餘忽略。
    """
    if isinstance(item, str):
        return item.strip()
    if isinstance(item, dict):
        action = item.get("action") or item.get("task") or ""
        owner = item.get("owner") or item.get("role") or ""
        due = item.get("due") or item.get("timeline") or ""
        kpi = item.get("kpi") or item.get("metric") or ""
        parts: List[str] = []
        if action:
            parts.append(action)
        if owner:
            parts.append(f"（責任：{owner}）")
        if due:
            parts.append(f"（時程：{due}）")
        if kpi:
            parts.append(f"（KPI：{kpi}）")
        return " ".join(parts).strip() or "(未提供內容)"
    return "(未提供內容)"


def parse_sections(ai_output: str) -> Tuple[str, List[str], List[str]]:
    """
    解析 JSON，抽出摘要/大綱/待辦。
    - 支援 todos 為字串或物件；最終輸出 list[str]
    - 忽略 self_evaluation 與 risks（若需要可擴充回傳）
    """
    try:
        data: Dict[str, Any] = json.loads(ai_output or "{}")
    except json.JSONDecodeError:
        data = {}

    summary = str(data.get("summary") or "").strip()
    outline_raw = data.get("outline") or []
    todos_raw = data.get("todos") or []

    # 清洗 outline
    outline: List[str] = []
    if isinstance(outline_raw, list):
        for x in outline_raw:
            if isinstance(x, str) and x.strip():
                outline.append(x.strip())

    # 清洗 todos
    todos: List[str] = []
    if isinstance(todos_raw, list):
        for item in todos_raw:
            s = _stringify_todo(item)
            if s:
                todos.append(s)

    return summary, outline, todos


# ✅ 提供主流程 batch_job.py 呼叫（簽名不變）
def run_analysis(text: str, preview: bool = False) -> Tuple[str, List[str], List[str]]:
    """
    主函式：輸入文字後回傳 (summary, outline, todos)
    - preview=True 時僅取前 500 字，快速出結果
    """
    if not text:
        raise ValueError("無效文字內容")

    clipped = text[:500] if preview else text
    ai_output = ai_analyze(clipped)
    if not ai_output:
        ai_output = "{}"

    return parse_sections(ai_output)


# 🔧 CLI 測試：直接生成 docx 方便本地驗收
def save_docx(path: str, summary: str, outline: List[str], todos: List[str]) -> None:
    doc = Document()
    doc.add_heading("會議摘要報告", level=0)

    doc.add_heading("摘要", level=1)
    doc.add_paragraph(summary if summary else "（無內容）")

    doc.add_heading("大綱", level=1)
    for item in (outline or ["（無內容）"]):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("待辦事項", level=1)
    for item in (todos or ["（無內容）"]):
        doc.add_paragraph(item, style="List Number")

    doc.save(path)


def main() -> None:
    # CLI 模式測試用：讀取 output/final/transcript_revised.txt
    input_path = "output/final/transcript_revised.txt"
    preview_mode = "--preview" in sys.argv

    if not os.path.exists(input_path):
        print(f"❌ 找不到輸入檔：{input_path}")
        return

    with open(input_path, "r", encoding="utf-8") as f:
        text = f.read()

    print(f"🤖 開始分析（preview={preview_mode}，model={OPENAI_MODEL}）")
    summary, outline, todos = run_analysis(text, preview=preview_mode)

    # 輸出結果
    os.makedirs("output/final", exist_ok=True)
    with open("output/final/meeting_summary.txt", "w", encoding="utf-8") as f:
        f.write("【摘要】\n" + summary + "\n\n【大綱】\n" + "\n".join(outline) +
                "\n\n【待辦事項】\n" + "\n".join(todos))

    with open("output/final/meeting_summary.json", "w", encoding="utf-8") as f:
        json.dump({"summary": summary, "outline": outline, "todos": todos},
                  f, ensure_ascii=False, indent=2)

    save_docx("output/final/meeting_summary.docx", summary, outline, todos)
    print("✅ 已輸出分析結果至 output/final/")


if __name__ == "__main__":
    main()
